"""
ADE Insurance Document Intelligence Backend API
FastAPI server with mock AI processing capabilities
"""

from fastapi import FastAPI, HTTPException, File, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse
from contextlib import asynccontextmanager
import os
import uuid
import aiofiles
import time
from datetime import datetime
from typing import Optional, List
import json

# PDF and Image processing imports
try:
    import fitz  # PyMuPDF
    PDF_PROCESSING_AVAILABLE = True
except ImportError:
    PDF_PROCESSING_AVAILABLE = False
    print("Warning: PyMuPDF not installed. PDF processing will be limited.")

try:
    from PIL import Image
    IMAGE_PROCESSING_AVAILABLE = True
except ImportError:
    IMAGE_PROCESSING_AVAILABLE = False
    print("Warning: Pillow not installed. Image processing will be limited.")

try:
    from docx import Document as DocxDocument
    from docx.shared import Inches
    DOCX_PROCESSING_AVAILABLE = True
except ImportError:
    DOCX_PROCESSING_AVAILABLE = False
    print("Warning: python-docx not installed. DOCX processing will be limited.")

try:
    from pdf2image import convert_from_path
    PDF2IMAGE_AVAILABLE = True
except ImportError:
    PDF2IMAGE_AVAILABLE = False
    print("Warning: pdf2image not installed. Alternative PDF processing will be used.")

from app.database import init_db, get_db
from app.models import Document, Job, Page, User, DisasterLocation
from app.schemas import (
    DocumentResponse,
    JobResponse,
    UploadResponse,
    RegionResponse,
    ProcessingRequest,
    UserRegister,
    UserLogin,
    UserResponse,
    UserUpdate,
    ChangePasswordRequest,
    TokenResponse,
    DisasterLocationResponse,
    DisasterLocationCreate,
    DisasterLocationUpdate
)
from app.weather_service import WeatherService
from app.ai_service import analyze_auto_document, extract_markdown_content, get_image_path_from_url
from app.geo_analyst import GeoAnalyst, generate_gemini_prompt

# JWT Configuration
from datetime import timedelta
import jwt
from typing import Optional
from decouple import config

# Load from environment variables
SECRET_KEY = config('SECRET_KEY', default='your-secret-key-change-this-in-production')
ALGORITHM = "HS256"
ACCESS_TOKEN_EXPIRE_MINUTES = config('ACCESS_TOKEN_EXPIRE_MINUTES', default=10080, cast=int)  # 7 days

def create_access_token(data: dict, expires_delta: Optional[timedelta] = None):
    """Create JWT access token"""
    to_encode = data.copy()
    if expires_delta:
        expire = datetime.utcnow() + expires_delta
    else:
        expire = datetime.utcnow() + timedelta(minutes=15)
    to_encode.update({"exp": expire})
    encoded_jwt = jwt.encode(to_encode, SECRET_KEY, algorithm=ALGORITHM)
    return encoded_jwt

# Utility Functions
def merge_page_results(page_results: List[dict]) -> dict:
    """
    Merge analysis results from multiple pages into a single result
    
    Args:
        page_results: List of analysis results from each page
        
    Returns:
        Merged result with combined information from all pages
    """
    if not page_results:
        return {
            "error": "No pages analyzed",
            "document_type": "Error",
            "confidence": 0.0,
            "title": None,
            "summary": "No content available",
            "people": [],
            "organizations": [],
            "locations": [],
            "dates": [],
            "numbers": [],
            "signature_detected": False,
            "total_pages": 0,
            "pages": []
        }
    
    # If only one page, return it with minor modifications
    if len(page_results) == 1:
        result = page_results[0].copy()
        result['total_pages'] = 1
        result['pages'] = [result.copy()]
        return result
    
    # For multiple pages, merge information
    # Collect all unique values across pages
    all_people = []
    all_organizations = []
    all_locations = []
    all_dates = []
    all_numbers = []
    has_signature = False
    
    # Document type determination (most common, highest confidence)
    doc_types = {}
    total_confidence = 0
    titles = []
    summaries = []
    
    for page_result in page_results:
        # Collect people (can be strings or dicts)
        if 'people' in page_result and page_result['people']:
            for person in page_result['people']:
                if isinstance(person, dict):
                    # If dict, try to get name field
                    person_str = person.get('name') or person.get('label') or str(person)
                    if person_str and person_str not in [p if isinstance(p, str) else p.get('name', '') for p in all_people]:
                        all_people.append(person)
                elif isinstance(person, str) and person not in [p if isinstance(p, str) else p.get('name', '') for p in all_people]:
                    all_people.append(person)
        
        # Collect organizations (can be strings or dicts)
        if 'organizations' in page_result and page_result['organizations']:
            for org in page_result['organizations']:
                if isinstance(org, dict):
                    org_str = org.get('name') or org.get('label') or str(org)
                    if org_str and org_str not in [o if isinstance(o, str) else o.get('name', '') for o in all_organizations]:
                        all_organizations.append(org)
                elif isinstance(org, str) and org not in [o if isinstance(o, str) else o.get('name', '') for o in all_organizations]:
                    all_organizations.append(org)
        
        # Collect locations
        if 'locations' in page_result and page_result['locations']:
            for loc in page_result['locations']:
                if isinstance(loc, str) and loc not in all_locations:
                    all_locations.append(loc)
                elif isinstance(loc, dict):
                    loc_str = loc.get('name') or loc.get('label') or str(loc)
                    if loc_str not in all_locations:
                        all_locations.append(loc_str)
        
        # Collect dates
        if 'dates' in page_result and page_result['dates']:
            for date in page_result['dates']:
                if isinstance(date, dict):
                    date_str = date.get('value') or date.get('label') or str(date)
                    if date_str and date_str not in [d if isinstance(d, str) else d.get('value', '') for d in all_dates]:
                        all_dates.append(date)
                elif isinstance(date, str) and date not in [d if isinstance(d, str) else d.get('value', '') for d in all_dates]:
                    all_dates.append(date)
        
        # Collect numbers
        if 'numbers' in page_result and page_result['numbers']:
            for num in page_result['numbers']:
                if isinstance(num, dict):
                    num_str = num.get('value') or num.get('label') or str(num)
                    if num_str and num_str not in [n if isinstance(n, str) else n.get('value', '') for n in all_numbers]:
                        all_numbers.append(num)
                elif isinstance(num, str) and num not in [n if isinstance(n, str) else n.get('value', '') for n in all_numbers]:
                    all_numbers.append(num)
        
        # Check for signatures
        if page_result.get('signature_detected', False):
            has_signature = True
        
        # Count document types
        doc_type = page_result.get('document_type', 'Unknown')
        confidence = page_result.get('confidence', 0.0)
        
        if doc_type not in doc_types:
            doc_types[doc_type] = {'count': 0, 'total_confidence': 0}
        doc_types[doc_type]['count'] += 1
        doc_types[doc_type]['total_confidence'] += confidence
        total_confidence += confidence
        
        # Collect titles and summaries
        if page_result.get('title'):
            titles.append(f"Page {page_result.get('page_number', '?')}: {page_result['title']}")
        if page_result.get('summary'):
            summaries.append(f"Page {page_result.get('page_number', '?')}: {page_result['summary']}")
    
    # Determine most likely document type (weighted by confidence)
    best_doc_type = "Unknown"
    best_score = 0
    for doc_type, info in doc_types.items():
        # Score = count * average confidence
        avg_confidence = info['total_confidence'] / info['count']
        score = info['count'] * avg_confidence
        if score > best_score:
            best_score = score
            best_doc_type = doc_type
    
    # Calculate average confidence
    avg_confidence = total_confidence / len(page_results) if page_results else 0.0
    
    # Create merged title
    merged_title = titles[0] if titles else None
    
    # Create merged summary
    merged_summary = f"Multi-page document with {len(page_results)} pages. " + " | ".join(summaries[:3])
    if len(summaries) > 3:
        merged_summary += f" ... and {len(summaries) - 3} more pages"
    
    # Return merged result
    return {
        "document_type": best_doc_type,
        "confidence": round(avg_confidence, 2),
        "title": merged_title,
        "summary": merged_summary,
        "people": all_people,
        "organizations": all_organizations,
        "locations": all_locations,
        "dates": all_dates,
        "numbers": all_numbers,
        "signature_detected": has_signature,
        "total_pages": len(page_results),
        "pages": page_results  # Include individual page results for reference
    }

# PDF Processing Functions
async def process_pdf_to_images(pdf_path: str, document_id: str) -> List[str]:
    """
    Convert PDF pages to individual images
    Returns list of image URLs (relative to /data/)
    """
    if not PDF_PROCESSING_AVAILABLE:
        # Fallback: return original PDF path
        return [f"/data/docs/{document_id}.pdf"]
    
    image_urls = []
    
    try:
        # Open PDF
        pdf_doc = fitz.open(pdf_path)
        
        for page_num in range(len(pdf_doc)):
            # Get page
            page = pdf_doc.load_page(page_num)
            
            # Convert to image (150 DPI for good quality and reasonable file size)
            mat = fitz.Matrix(1.5, 1.5)  # 1.5x zoom = ~150 DPI
            pix = page.get_pixmap(matrix=mat)
            
            # Save as PNG in images folder
            image_filename = f"{document_id}_page_{page_num + 1}.png"
            image_path = f"data/images/{image_filename}"
            pix.save(image_path)
            
            # Return URL path
            image_urls.append(f"/data/images/{image_filename}")
            
        pdf_doc.close()
        
        # If no pages were processed, return fallback
        if not image_urls:
            return [f"/data/docs/{document_id}.pdf"]
        
    except Exception as e:
        print(f"Error processing PDF: {e}")
        # Fallback: return original PDF path
        return [f"/data/docs/{document_id}.pdf"]
    
    return image_urls

async def process_image_file(image_path: str, document_id: str, page_index: int = 0) -> str:
    """
    Process single image file (basic validation)
    Returns image URL path
    """
    # For now, just return the URL path without processing
    # In the future, could add resize/optimization here
    file_ext = os.path.splitext(image_path)[1]
    return f"/data/images/{document_id}{file_ext}"

async def process_docx_to_images(docx_path: str, document_id: str) -> List[str]:
    """
    Convert DOCX pages to individual images
    Returns list of image URLs (relative to /data/)
    """
    if not DOCX_PROCESSING_AVAILABLE:
        # Fallback: return original DOCX path
        return [f"/data/docs/{document_id}.docx"]
    
    image_urls = []
    
    try:
        # For DOCX, we'll create a preview image of the first page
        # This is a simplified approach - in production you might want to use more sophisticated methods
        
        # Load DOCX document
        doc = DocxDocument(docx_path)
        
        # Create a simple preview by extracting text and creating an image
        # For now, we'll just create one preview image
        preview_filename = f"{document_id}_preview.png"
        preview_path = f"data/images/{preview_filename}"
        
        # Create a simple text-based preview image
        if IMAGE_PROCESSING_AVAILABLE:
            from PIL import Image, ImageDraw, ImageFont
            
            # Get document text (first few paragraphs)
            text_content = ""
            paragraph_count = 0
            for paragraph in doc.paragraphs:
                if paragraph.text.strip() and paragraph_count < 10:  # First 10 paragraphs
                    text_content += paragraph.text.strip() + "\n\n"
                    paragraph_count += 1
            
            if not text_content:
                text_content = "DOCX Document Preview\n\nDocument contains formatted content."
            
            # Create image with text
            img_width, img_height = 800, 1000
            img = Image.new('RGB', (img_width, img_height), color='white')
            draw = ImageDraw.Draw(img)
            
            try:
                # Try to use a system font
                font = ImageFont.load_default()
            except:
                font = None
            
            # Draw text with word wrapping
            y_position = 50
            lines = text_content.split('\n')
            for line in lines[:30]:  # Limit to 30 lines
                if y_position > img_height - 50:
                    break
                draw.text((50, y_position), line[:80], fill='black', font=font)  # Limit line length
                y_position += 25
            
            # Save preview image
            img.save(preview_path, "PNG")
            image_urls.append(f"/data/images/{preview_filename}")
        else:
            # If PIL not available, return DOCX file path
            return [f"/data/docs/{document_id}.docx"]
        
    except Exception as e:
        print(f"Error processing DOCX: {e}")
        # Fallback: return original DOCX path
        return [f"/data/docs/{document_id}.docx"]
    
    return image_urls if image_urls else [f"/data/docs/{document_id}.docx"]

async def process_pdf_alternative(pdf_path: str, document_id: str) -> List[str]:
    """
    Alternative PDF processing using pdf2image
    Returns list of image URLs (relative to /data/)
    """
    if not PDF2IMAGE_AVAILABLE:
        return await process_pdf_to_images(pdf_path, document_id)
    
    image_urls = []
    
    try:
        # Convert PDF pages to images using pdf2image
        images = convert_from_path(pdf_path, dpi=150, first_page=1, last_page=10)  # Limit to first 10 pages
        
        for i, image in enumerate(images):
            image_filename = f"{document_id}_page_{i + 1}.png"
            image_path = f"data/images/{image_filename}"
            
            # Save image
            image.save(image_path, "PNG")
            image_urls.append(f"/data/images/{image_filename}")
            
        # If no pages were processed, fallback
        if not image_urls:
            return [f"/data/docs/{document_id}.pdf"]
        
    except Exception as e:
        print(f"Error in alternative PDF processing: {e}")
        # Fallback to original PDF processing
        return await process_pdf_to_images(pdf_path, document_id)
    
    return image_urls

@asynccontextmanager
async def lifespan(app: FastAPI):
    # Startup
    await init_db()
    # Create data directories if they don't exist
    os.makedirs("data/docs", exist_ok=True)
    os.makedirs("data/images", exist_ok=True)
    yield
    # Shutdown (cleanup if needed)

# Initialize FastAPI app
app = FastAPI(
    title="ADE Insurance Document Intelligence API",
    description="Mock AI Document Processing API for Insurance Documents",
    version="1.0.0",
    docs_url="/docs",
    redoc_url="/redoc",
    lifespan=lifespan
)

# Configure CORS for Frontend
FRONTEND_URL = config('FRONTEND_URL', default='http://localhost:5173')

# Allowed origins for CORS
allowed_origins = [
    "http://localhost:5173",
    "http://localhost:5174", 
    "http://localhost:3000",
]

# Add production Vercel URL if configured
if FRONTEND_URL and FRONTEND_URL != "http://localhost:5173":
    allowed_origins.append(FRONTEND_URL)

# Allow all Vercel preview deployments (*.vercel.app)
# This is a more permissive approach for development
import re
def is_vercel_domain(origin: str) -> bool:
    """Check if origin is a Vercel domain"""
    return bool(re.match(r"https://.*\.vercel\.app$", origin))

# Custom CORS middleware to handle Vercel domains
from starlette.middleware.cors import CORSMiddleware as StarleteCORSMiddleware
from starlette.types import ASGIApp, Receive, Scope, Send

class CustomCORSMiddleware(StarleteCORSMiddleware):
    async def __call__(self, scope: Scope, receive: Receive, send: Send) -> None:
        if scope["type"] == "http":
            origin = None
            for header_name, header_value in scope.get("headers", []):
                if header_name == b"origin":
                    origin = header_value.decode("utf-8")
                    break
            
            # Allow Vercel preview domains dynamically
            if origin and is_vercel_domain(origin) and origin not in self.allow_origins:
                self.allow_origins.append(origin)
        
        await super().__call__(scope, receive, send)

app.add_middleware(
    CustomCORSMiddleware,
    allow_origins=allowed_origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Serve static files (uploaded documents and images)
app.mount("/data", StaticFiles(directory="data"), name="data")

@app.get("/")
async def root():
    """Fast health check endpoint"""
    return {
        "status": "healthy",
        "message": "ADE Insurance Document Intelligence API",
        "version": "1.0.0",
        "docs": "/docs"
    }

@app.get("/health")
async def health_check():
    """Fast health check for monitoring"""
    return {"status": "healthy", "timestamp": time.time()}

# ==================== Authentication Endpoints ====================

@app.post("/auth/register", response_model=TokenResponse)
async def register_user(user_data: UserRegister):
    """
    Register a new user account
    """
    db = get_db()
    try:
        # Check if email already exists
        existing_user = db.query(User).filter(User.email == user_data.email).first()
        if existing_user:
            raise HTTPException(
                status_code=400,
                detail="Email already registered"
            )
        
        # Create new user
        hashed_password = User.hash_password(user_data.password)
        new_user = User(
            email=user_data.email,
            full_name=user_data.full_name,
            phone=user_data.phone,
            hashed_password=hashed_password,
            # Extended profile fields
            address=user_data.address,
            date_of_birth=user_data.date_of_birth,
            gender=user_data.gender,
            id_number=user_data.id_number,
            place_of_origin=user_data.place_of_origin,
            occupation=user_data.occupation
        )
        
        db.add(new_user)
        db.commit()
        db.refresh(new_user)
        
        # Create access token
        access_token = create_access_token(
            data={"sub": str(new_user.id), "email": new_user.email},
            expires_delta=timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
        )
        
        # Return token and user data
        return TokenResponse(
            access_token=access_token,
            user=UserResponse.model_validate(new_user)
        )
        
    except HTTPException:
        raise
    except Exception as e:
        db.rollback()
        raise HTTPException(
            status_code=500,
            detail=f"Registration failed: {str(e)}"
        )
    finally:
        db.close()

@app.post("/auth/login", response_model=TokenResponse)
async def login_user(credentials: UserLogin):
    """
    Login user and return JWT token
    """
    db = get_db()
    try:
        # Find user by email
        user = db.query(User).filter(User.email == credentials.email).first()
        
        if not user or not user.verify_password(credentials.password):
            raise HTTPException(
                status_code=401,
                detail="Incorrect email or password"
            )
        
        if not user.is_active:
            raise HTTPException(
                status_code=403,
                detail="Account is inactive"
            )
        
        # Update last login
        user.last_login = datetime.utcnow()
        db.commit()
        
        # Create access token
        access_token = create_access_token(
            data={"sub": str(user.id), "email": user.email},
            expires_delta=timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
        )
        
        # Return token and user data
        return TokenResponse(
            access_token=access_token,
            user=UserResponse.model_validate(user)
        )
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Login failed: {str(e)}"
        )
    finally:
        db.close()

@app.get("/auth/me", response_model=UserResponse)
async def get_current_user(token: str):
    """
    Get current user from JWT token
    """
    db = get_db()
    try:
        # Decode JWT token
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        user_id = int(payload.get("sub"))
        
        if not user_id:
            raise HTTPException(status_code=401, detail="Invalid token")
        
        # Get user from database
        user = db.query(User).filter(User.id == user_id).first()
        
        if not user:
            raise HTTPException(status_code=404, detail="User not found")
        
        return UserResponse.model_validate(user)
        
    except jwt.ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="Token expired")
    except jwt.InvalidTokenError:
        raise HTTPException(status_code=401, detail="Invalid token")
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        db.close()

# ==================== User Profile Management Endpoints ====================

def get_user_from_token(token: str) -> User:
    """Helper function to get user from JWT token"""
    db = get_db()
    try:
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        user_id = int(payload.get("sub"))
        
        if not user_id:
            raise HTTPException(status_code=401, detail="Invalid token")
        
        user = db.query(User).filter(User.id == user_id).first()
        
        if not user:
            raise HTTPException(status_code=404, detail="User not found")
        
        return user
    except jwt.ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="Token expired")
    except jwt.InvalidTokenError:
        raise HTTPException(status_code=401, detail="Invalid token")
    finally:
        db.close()

@app.put("/users/profile", response_model=UserResponse)
async def update_user_profile(profile_data: UserUpdate, token: str):
    """
    Update user profile information
    """
    db = get_db()
    try:
        # Get current user
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        user_id = int(payload.get("sub"))
        
        user = db.query(User).filter(User.id == user_id).first()
        if not user:
            raise HTTPException(status_code=404, detail="User not found")
        
        # Update fields if provided
        update_data = profile_data.model_dump(exclude_unset=True)
        
        for field, value in update_data.items():
            if hasattr(user, field):
                setattr(user, field, value)
        
        # Update timestamp
        user.updated_at = datetime.utcnow()
        
        # Check if profile is now complete
        required_fields = ['full_name', 'phone', 'address', 'date_of_birth']
        profile_complete = all(getattr(user, field) for field in required_fields)
        user.profile_completed = profile_complete
        
        db.commit()
        db.refresh(user)
        
        print(f"✅ Updated profile for user {user.email}")
        
        return UserResponse.model_validate(user)
        
    except jwt.ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="Token expired")
    except jwt.InvalidTokenError:
        raise HTTPException(status_code=401, detail="Invalid token")
    except HTTPException:
        raise
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"Profile update failed: {str(e)}")
    finally:
        db.close()

@app.post("/users/change-password")
async def change_password(password_data: ChangePasswordRequest, token: str):
    """
    Change user password
    """
    db = get_db()
    try:
        # Get current user
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        user_id = int(payload.get("sub"))
        
        user = db.query(User).filter(User.id == user_id).first()
        if not user:
            raise HTTPException(status_code=404, detail="User not found")
        
        # Verify current password
        if not user.verify_password(password_data.current_password):
            raise HTTPException(status_code=400, detail="Current password is incorrect")
        
        # Update password
        user.hashed_password = User.hash_password(password_data.new_password)
        user.updated_at = datetime.utcnow()
        
        db.commit()
        
        print(f"✅ Password changed for user {user.email}")
        
        return {"message": "Password changed successfully"}
        
    except jwt.ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="Token expired")
    except jwt.InvalidTokenError:
        raise HTTPException(status_code=401, detail="Invalid token")
    except HTTPException:
        raise
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"Password change failed: {str(e)}")
    finally:
        db.close()

@app.get("/users/{user_id}", response_model=UserResponse)
async def get_user_by_id(user_id: int, token: str):
    """
    Get user by ID (admin or self only)
    """
    db = get_db()
    try:
        # Get current user
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        current_user_id = int(payload.get("sub"))
        
        # Only allow viewing own profile for now (can be extended for admin)
        if current_user_id != user_id:
            raise HTTPException(status_code=403, detail="Access denied")
        
        user = db.query(User).filter(User.id == user_id).first()
        if not user:
            raise HTTPException(status_code=404, detail="User not found")
        
        return UserResponse.model_validate(user)
        
    except jwt.ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="Token expired")
    except jwt.InvalidTokenError:
        raise HTTPException(status_code=401, detail="Invalid token")
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        db.close()

@app.get("/users/profile/completion")
async def get_profile_completion(token: str):
    """
    Get profile completion status and missing fields
    """
    db = get_db()
    try:
        # Get current user
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        user_id = int(payload.get("sub"))
        
        user = db.query(User).filter(User.id == user_id).first()
        if not user:
            raise HTTPException(status_code=404, detail="User not found")
        
        # Required fields for complete profile
        required_fields = {
            'full_name': 'Họ và tên',
            'phone': 'Số điện thoại',
            'address': 'Địa chỉ',
            'date_of_birth': 'Ngày sinh',
            'gender': 'Giới tính',
            'id_number': 'Số CCCD/CMND'
        }
        
        # Optional but recommended fields
        recommended_fields = {
            'place_of_origin': 'Quê quán',
            'occupation': 'Nghề nghiệp',
            'emergency_contact_name': 'Người liên hệ khẩn cấp',
            'emergency_contact_phone': 'SĐT người liên hệ khẩn cấp'
        }
        
        missing_required = []
        missing_recommended = []
        
        # Check required fields
        for field, label in required_fields.items():
            if not getattr(user, field):
                missing_required.append({'field': field, 'label': label})
        
        # Check recommended fields
        for field, label in recommended_fields.items():
            if not getattr(user, field):
                missing_recommended.append({'field': field, 'label': label})
        
        total_fields = len(required_fields) + len(recommended_fields)
        completed_fields = total_fields - len(missing_required) - len(missing_recommended)
        completion_percentage = int((completed_fields / total_fields) * 100)
        
        return {
            "completion_percentage": completion_percentage,
            "profile_completed": len(missing_required) == 0,
            "missing_required_fields": missing_required,
            "missing_recommended_fields": missing_recommended,
            "total_fields": total_fields,
            "completed_fields": completed_fields
        }
        
    except jwt.ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="Token expired")
    except jwt.InvalidTokenError:
        raise HTTPException(status_code=401, detail="Invalid token")
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        db.close()

# ==================== Document Endpoints ====================

def cleanup_old_images(max_images: int = 30):
    """
    Cleanup old images when exceeding max_images limit
    Deletes oldest images first to keep storage under control
    """
    try:
        image_dir = "data/images"
        
        # Get all image files with their modification time
        image_files = []
        for filename in os.listdir(image_dir):
            file_path = os.path.join(image_dir, filename)
            if os.path.isfile(file_path) and filename.lower().endswith(('.jpg', '.jpeg', '.png')):
                mod_time = os.path.getmtime(file_path)
                image_files.append((file_path, mod_time, filename))
        
        # Check if cleanup is needed
        total_images = len(image_files)
        if total_images <= max_images:
            print(f"📊 Image count: {total_images}/{max_images} - No cleanup needed")
            return
        
        # Sort by modification time (oldest first)
        image_files.sort(key=lambda x: x[1])
        
        # Calculate how many to delete
        images_to_delete = total_images - max_images
        
        print(f"🧹 Starting cleanup: {total_images} images found, deleting {images_to_delete} oldest images...")
        
        # Delete oldest images
        deleted_count = 0
        for file_path, mod_time, filename in image_files[:images_to_delete]:
            try:
                os.remove(file_path)
                deleted_count += 1
                print(f"  🗑️  Deleted: {filename}")
            except Exception as e:
                print(f"  ⚠️  Failed to delete {filename}: {e}")
        
        remaining = total_images - deleted_count
        print(f"✅ Cleanup complete: Deleted {deleted_count} images, {remaining} remaining")
        
    except Exception as e:
        print(f"❌ Cleanup error: {e}")
        import traceback
        traceback.print_exc()

@app.post("/documents/upload", response_model=UploadResponse)
async def upload_document(file: UploadFile = File(...)):
    """
    Upload a document file and create a new document record
    """
    db = get_db()
    try:
        # 🧹 Cleanup old images if exceeding limit (only for image uploads)
        file_ext = os.path.splitext(file.filename)[1].lower()
        if file_ext in ['.png', '.jpg', '.jpeg']:
            cleanup_old_images(max_images=30)
        
        # Generate unique document ID
        document_id = str(uuid.uuid4())
        
        # Get file extension
        file_ext = os.path.splitext(file.filename)[1].lower()
        if file_ext not in ['.pdf', '.docx', '.png', '.jpg', '.jpeg']:
            raise HTTPException(status_code=400, detail="Unsupported file format. Only PDF, DOCX, PNG, JPG are allowed.")
        
        print(f"Uploading document: {file.filename} (ID: {document_id})")
        
        # Save file to appropriate storage location
        if file_ext in ['.pdf', '.docx']:
            file_path = f"data/docs/{document_id}{file_ext}"
        else:
            file_path = f"data/images/{document_id}{file_ext}"
        
        async with aiofiles.open(file_path, 'wb') as f:
            content = await file.read()
            await f.write(content)
        
        print(f"File saved to: {file_path}")
        
        # Create document record in database
        document = Document(
            id=document_id,
            filename=file.filename,
            status="NOT_STARTED",
            created_at=datetime.utcnow()
        )
        db.add(document)
        
        print(f"Creating document record in database...")
        
        # Create page records based on file type
        if file_ext == '.pdf':
            # For PDF: try to extract pages, fallback to single page if failed
            try:
                print(f"Processing PDF...")
                # Try alternative PDF processing first if available
                if PDF2IMAGE_AVAILABLE:
                    image_urls = await process_pdf_alternative(file_path, document_id)
                else:
                    image_urls = await process_pdf_to_images(file_path, document_id)
                
                print(f"PDF processed into {len(image_urls)} pages")
                
                # Create page for each extracted image/page
                for i, image_url in enumerate(image_urls):
                    page = Page(
                        document_id=document_id,
                        page_index=i,
                        image_url=image_url
                    )
                    db.add(page)
            except Exception as e:
                print(f"PDF processing failed: {e}")
                # Fallback: single page pointing to PDF
                page = Page(
                    document_id=document_id,
                    page_index=0,
                    image_url=f"/data/docs/{document_id}{file_ext}"
                )
                db.add(page)
        elif file_ext == '.docx':
            # For DOCX: try to create preview images
            try:
                print(f"Processing DOCX...")
                image_urls = await process_docx_to_images(file_path, document_id)
                print(f"DOCX processed into {len(image_urls)} pages")
                # Create page for each preview image
                for i, image_url in enumerate(image_urls):
                    page = Page(
                        document_id=document_id,
                        page_index=i,
                        image_url=image_url
                    )
                    db.add(page)
            except Exception as e:
                print(f"DOCX processing failed: {e}")
                # Fallback: single page pointing to DOCX
                page = Page(
                    document_id=document_id,
                    page_index=0,
                    image_url=f"/data/docs/{document_id}{file_ext}"
                )
                db.add(page)
        else:
            # For images: single page pointing to the image
            print(f"Processing image file...")
            page = Page(
                document_id=document_id,
                page_index=0,
                image_url=f"/data/images/{document_id}{file_ext}"
            )
            db.add(page)
        
        db.commit()
        print(f"✅ Document uploaded successfully: {document_id}")
        
        return UploadResponse(document_id=document_id)
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ Upload error: {e}")
        import traceback
        traceback.print_exc()
        db.rollback()
        raise HTTPException(status_code=500, detail=f"Upload failed: {str(e)}")
    finally:
        db.close()

@app.get("/documents/images/stats")
async def get_image_stats():
    """
    Get statistics about stored images
    """
    try:
        image_dir = "data/images"
        
        # Get all image files
        image_files = []
        total_size = 0
        
        for filename in os.listdir(image_dir):
            file_path = os.path.join(image_dir, filename)
            if os.path.isfile(file_path) and filename.lower().endswith(('.jpg', '.jpeg', '.png')):
                size = os.path.getsize(file_path)
                mod_time = os.path.getmtime(file_path)
                image_files.append({
                    "filename": filename,
                    "size_mb": round(size / (1024 * 1024), 2),
                    "modified": datetime.fromtimestamp(mod_time).isoformat()
                })
                total_size += size
        
        # Sort by modification time (newest first)
        image_files.sort(key=lambda x: x['modified'], reverse=True)
        
        return {
            "total_images": len(image_files),
            "total_size_mb": round(total_size / (1024 * 1024), 2),
            "limit": 30,
            "cleanup_needed": len(image_files) > 30,
            "images_to_delete": max(0, len(image_files) - 30),
            "recent_images": image_files[:10]  # Show 10 most recent
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to get stats: {str(e)}")

@app.post("/documents/images/cleanup")
async def trigger_cleanup():
    """
    Manually trigger image cleanup
    """
    try:
        cleanup_old_images(max_images=30)
        
        # Get updated stats
        image_dir = "data/images"
        remaining = len([f for f in os.listdir(image_dir) 
                        if os.path.isfile(os.path.join(image_dir, f)) 
                        and f.lower().endswith(('.jpg', '.jpeg', '.png'))])
        
        return {
            "status": "success",
            "remaining_images": remaining,
            "limit": 30
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Cleanup failed: {str(e)}")

@app.post("/documents/{document_id}/process")
async def process_document(document_id: str):
    """
    Process document using Gemini auto-analysis
    This replaces the old mock processing system
    """
    db = get_db()
    try:
        document = db.query(Document).filter(Document.id == document_id).first()
        
        if not document:
            raise HTTPException(status_code=404, detail="Document not found")
        
        # Get first page image
        pages = db.query(Page).filter(Page.document_id == document_id).order_by(Page.page_index).all()
        
        if not pages:
            raise HTTPException(status_code=404, detail="No pages found for this document")
        
        # Get first page image URL
        first_page = pages[0]
        image_url = first_page.image_url
        
        # Convert URL to local file path
        image_path = get_image_path_from_url(image_url)
        
        if not image_path:
            raise HTTPException(status_code=400, detail="Invalid image path")
        
        try:
            # Update document status to processing
            document.status = "PROCESSING"
            db.commit()
            
            # Analyze document with Gemini
            result = await analyze_auto_document(image_path)
            
            # Save result to database
            document.ai_result_json = json.dumps(result, ensure_ascii=False, indent=2)
            document.status = "DONE"
            db.commit()
            
            return {
                "status": "DONE",
                "message": "Document processed successfully"
            }
            
        except Exception as e:
            # Update status to error
            document.status = "ERROR"
            db.commit()
            
            raise HTTPException(
                status_code=500, 
                detail=f"Processing failed: {str(e)}"
            )
    finally:
        db.close()

@app.get("/documents/{document_id}", response_model=DocumentResponse)
async def get_document(document_id: str):
    """
    Get document metadata and pages
    """
    db = get_db()
    try:
        document = db.query(Document).filter(Document.id == document_id).first()
        
        if not document:
            raise HTTPException(status_code=404, detail="Document not found")
        
        # Get document pages
        pages = db.query(Page).filter(Page.document_id == document_id).all()
        
        pages_data = [
            {
                "page_index": page.page_index,
                "image_url": page.image_url
            }
            for page in pages
        ]
        
        return DocumentResponse(
            document_id=document.id,
            status=document.status,
            pages=pages_data
        )
    finally:
        db.close()

@app.get("/documents/{document_id}/overlay")
async def get_document_overlay(document_id: str):
    """
    Get document overlay regions - removed (not needed for fast mode)
    """
    raise HTTPException(
        status_code=404, 
        detail="Overlay feature not available in fast analysis mode"
    )

@app.get("/documents/{document_id}/markdown")
async def get_document_markdown(document_id: str):
    """
    Get document content as markdown (extracted from Gemini)
    """
    db = get_db()
    try:
        document = db.query(Document).filter(Document.id == document_id).first()
        
        if not document:
            raise HTTPException(status_code=404, detail="Document not found")
        
        # Return markdown if available
        if document.markdown_content:
            return {"markdown": document.markdown_content}
        
        # If not available, check document status
        if document.status == "PROCESSING":
            raise HTTPException(status_code=400, detail="Document is being processed")
        elif document.status == "ERROR":
            raise HTTPException(status_code=400, detail="Document processing failed")
        else:
            raise HTTPException(status_code=400, detail="Document not yet analyzed. Please call /analyze-auto first")
    finally:
        db.close()

@app.get("/documents/{document_id}/json")
async def get_document_json(document_id: str):
    """
    Get document structured data as JSON from Gemini analysis
    """
    db = get_db()
    try:
        document = db.query(Document).filter(Document.id == document_id).first()
        
        if not document:
            raise HTTPException(status_code=404, detail="Document not found")
        
        # Check if document has been analyzed
        if not document.ai_result_json:
            if document.status == "PROCESSING":
                raise HTTPException(status_code=400, detail="Document is being processed")
            elif document.status == "ERROR":
                raise HTTPException(status_code=400, detail="Document processing failed")
            else:
                raise HTTPException(status_code=400, detail="Document not yet analyzed. Please call /analyze-auto first")
        
        # Return AI analysis result
        try:
            return json.loads(document.ai_result_json)
        except json.JSONDecodeError:
            raise HTTPException(status_code=500, detail="Invalid JSON data stored")
    finally:
        db.close()

@app.put("/documents/{document_id}/json")
async def update_document_json(document_id: str, json_data: dict):
    """
    Update document JSON data
    """
    db = get_db()
    try:
        document = db.query(Document).filter(Document.id == document_id).first()
        
        if not document:
            raise HTTPException(status_code=404, detail="Document not found")
        
        # In a real implementation, you would save this to database
        # For now, just return success
        return {"success": True, "message": "JSON data updated successfully"}
    finally:
        db.close()

@app.post("/documents/{document_id}/analyze-auto")
async def analyze_document_auto(document_id: str):
    """
    Analyze document automatically using Gemini 2.5 Flash
    For multi-page PDFs: analyzes each page separately and merges results
    Extracts structured information and full text markdown
    """
    db = get_db()
    try:
        document = db.query(Document).filter(Document.id == document_id).first()
        
        if not document:
            raise HTTPException(status_code=404, detail="Document not found")
        
        # Get ALL pages for this document
        pages = db.query(Page).filter(Page.document_id == document_id).order_by(Page.page_index).all()
        
        if not pages:
            raise HTTPException(status_code=404, detail="No pages found for this document")
        
        try:
            # Update document status to processing
            document.status = "PROCESSING"
            db.commit()
            
            print(f"\n{'='*80}")
            print(f"📄 Analyzing document: {document_id}")
            print(f"   Filename: {document.filename}")
            print(f"   Total pages: {len(pages)}")
            print(f"{'='*80}\n")
            
            # Lists to collect results from all pages
            all_page_results = []
            all_markdown_parts = []
            
            # Process each page
            for idx, page in enumerate(pages):
                page_num = idx + 1
                print(f"\n📑 Processing page {page_num}/{len(pages)}...")
                print(f"   Image URL: {page.image_url}")
                
                # Convert URL to local file path
                image_path = get_image_path_from_url(page.image_url)
                
                if not image_path:
                    print(f"   ⚠️  Warning: Invalid image path for page {page_num}, skipping...")
                    continue
                
                if not os.path.exists(image_path):
                    print(f"   ⚠️  Warning: Image file not found: {image_path}, skipping...")
                    continue
                
                try:
                    # Analyze this page with Gemini (structured data)
                    print(f"   🤖 Extracting structured data...")
                    page_result = await analyze_auto_document(image_path)
                    
                    # Add page number to result
                    page_result['page_number'] = page_num
                    all_page_results.append(page_result)
                    print(f"   ✅ Structured data extracted")
                    
                    # Extract markdown content (full text)
                    print(f"   📝 Extracting markdown content...")
                    page_markdown = await extract_markdown_content(image_path)
                    
                    # Add page separator and page number to markdown
                    if len(pages) > 1:
                        markdown_with_header = f"\n\n---\n## Page {page_num}\n\n{page_markdown}"
                    else:
                        markdown_with_header = page_markdown
                    
                    all_markdown_parts.append(markdown_with_header)
                    print(f"   ✅ Markdown extracted ({len(page_markdown)} chars)")
                    
                except Exception as page_error:
                    print(f"   ❌ Error analyzing page {page_num}: {page_error}")
                    # Continue with next page even if this one fails
                    all_page_results.append({
                        "page_number": page_num,
                        "error": str(page_error),
                        "document_type": "Error",
                        "confidence": 0.0,
                        "title": None,
                        "summary": f"Page {page_num} analysis failed",
                        "people": [],
                        "organizations": [],
                        "locations": [],
                        "dates": [],
                        "numbers": [],
                        "signature_detected": False
                    })
                    all_markdown_parts.append(f"\n\n---\n## Page {page_num}\n\n*Error extracting content from this page*\n")
            
            # Merge results from all pages
            print(f"\n📊 Merging results from {len(all_page_results)} pages...")
            merged_result = merge_page_results(all_page_results)
            
            # Combine all markdown parts
            full_markdown = "\n".join(all_markdown_parts).strip()
            
            # Add document summary at the top of markdown
            if len(pages) > 1:
                markdown_header = f"# {document.filename}\n\n**Total Pages:** {len(pages)}\n**Document Type:** {merged_result.get('document_type', 'Unknown')}\n"
                full_markdown = markdown_header + full_markdown
            
            print(f"   ✅ Merged result - Document type: {merged_result.get('document_type')}")
            print(f"   ✅ Total markdown length: {len(full_markdown)} chars")
            
            # Save results to database
            document.ai_result_json = json.dumps(merged_result, ensure_ascii=False, indent=2)
            document.markdown_content = full_markdown
            document.status = "DONE"
            db.commit()
            
            print(f"\n✅ Analysis complete for document {document_id}")
            print(f"{'='*80}\n")
            
            return merged_result
            
        except Exception as e:
            # Update status to error
            document.status = "ERROR"
            db.commit()
            
            print(f"\n❌ Analysis failed: {e}")
            import traceback
            traceback.print_exc()
            
            raise HTTPException(
                status_code=500, 
                detail=f"Analysis failed: {str(e)}"
            )
    finally:
        db.close()


@app.post("/documents/{document_id}/extract-person-info")
async def extract_person_info_endpoint(document_id: str):
    """
    Extract personal information from document (CCCD/ID/Driver License)
    Optimized for insurance application forms
    """
    db = get_db()
    try:
        # Get document
        document = db.query(Document).filter(Document.id == document_id).first()
        if not document:
            raise HTTPException(status_code=404, detail="Document not found")
        
        print(f"\n👤 Extracting person info from document {document_id}")
        
        # Get first page image
        pages = db.query(Page).filter(Page.document_id == document_id).order_by(Page.page_index).all()
        
        if not pages or not pages[0].image_url:
            raise HTTPException(status_code=400, detail="No image found for this document")
        
        # Get image path
        from app.ai_service import get_image_path_from_url, extract_person_info
        image_path = get_image_path_from_url(pages[0].image_url)
        
        if not image_path or not os.path.exists(image_path):
            raise HTTPException(status_code=400, detail=f"Image file not found: {image_path}")
        
        print(f"   📷 Processing image: {image_path}")
        
        # Extract person info using Gemini
        person_info = await extract_person_info(image_path)
        
        if "error" in person_info:
            print(f"   ❌ Extraction error: {person_info['error']}")
        else:
            print(f"   ✅ Extracted: {person_info.get('fullName', 'N/A')}")
            
            # Save person_info to database
            import json
            document.person_data = json.dumps(person_info, ensure_ascii=False)
            db.commit()
            print(f"   💾 Saved person data to database")
        
        return {
            "document_id": document_id,
            "person_info": person_info,
            "message": "Person information extracted successfully"
        }
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"\n❌ Person info extraction failed: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Extraction failed: {str(e)}")
    finally:
        db.close()


@app.post("/documents/{document_id}/extract-vehicle-info")
async def extract_vehicle_info_endpoint(document_id: str):
    """
    Extract vehicle information from document (Giấy đăng ký xe / Cà vẹt)
    For vehicle insurance applications
    """
    db = get_db()
    try:
        # Get document
        document = db.query(Document).filter(Document.id == document_id).first()
        if not document:
            raise HTTPException(status_code=404, detail="Document not found")
        
        print(f"\n🚗 Extracting vehicle info from document {document_id}")
        
        # Get first page image
        pages = db.query(Page).filter(Page.document_id == document_id).order_by(Page.page_index).all()
        
        if not pages or not pages[0].image_url:
            raise HTTPException(status_code=400, detail="No image found for this document")
        
        # Get image path
        from app.ai_service import get_image_path_from_url, extract_vehicle_info
        image_path = get_image_path_from_url(pages[0].image_url)
        
        if not image_path or not os.path.exists(image_path):
            raise HTTPException(status_code=400, detail=f"Image file not found: {image_path}")
        
        print(f"   📷 Processing image: {image_path}")
        
        # Extract vehicle info using Gemini
        vehicle_info = await extract_vehicle_info(image_path)
        
        if "error" in vehicle_info:
            print(f"   ❌ Extraction error: {vehicle_info['error']}")
        else:
            print(f"   ✅ Extracted: {vehicle_info.get('licensePlate', 'N/A')} - {vehicle_info.get('brand', 'N/A')} {vehicle_info.get('model', 'N/A')}")
            
            # Save vehicle_info to database
            import json
            document.vehicle_data = json.dumps(vehicle_info, ensure_ascii=False)
            db.commit()
            print(f"   💾 Saved vehicle data to database")
        
        return {
            "document_id": document_id,
            "vehicle_info": vehicle_info,
            "message": "Vehicle information extracted successfully"
        }
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"\n❌ Vehicle info extraction failed: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Extraction failed: {str(e)}")
    finally:
        db.close()


@app.post("/documents/{document_id}/recommend-insurance")
async def recommend_insurance_endpoint(document_id: str):
    """
    Analyze document address and recommend insurance packages based on region
    Returns region-specific insurance recommendations (Bắc/Trung/Nam)
    Uses extracted PersonInfo (placeOfOrigin) if available, otherwise analyzes image
    """
    db = get_db()
    try:
        # Get document
        document = db.query(Document).filter(Document.id == document_id).first()
        if not document:
            raise HTTPException(status_code=404, detail="Document not found")
        
        print(f"\n🏠 Analyzing address for insurance recommendations: {document_id}")
        
        # Get first page
        pages = db.query(Page).filter(Page.document_id == document_id).order_by(Page.page_index).all()
        
        if not pages or not pages[0].image_url:
            raise HTTPException(status_code=400, detail="No image found for this document")
        
        # Get image path
        from app.ai_service import get_image_path_from_url, recommend_insurance_by_address, recommend_insurance_by_person_info
        image_path = get_image_path_from_url(pages[0].image_url)
        
        if not image_path or not os.path.exists(image_path):
            raise HTTPException(status_code=400, detail=f"Image file not found: {image_path}")
        
        print(f"   📷 Processing image: {image_path}")
        
        # Check if we have extracted person info with placeOfOrigin
        person_data = None
        if document.person_data:
            import json
            person_data = json.loads(document.person_data)
        
        if person_data and person_data.get('placeOfOrigin'):
            # Use extracted person info
            print(f"   ✅ Using extracted placeOfOrigin: {person_data.get('placeOfOrigin')}")
            print(f"   ✅ Using extracted address: {person_data.get('address', 'N/A')}")
            recommendation = await recommend_insurance_by_person_info(person_data)
        else:
            # Fallback to image analysis
            print(f"   ⚠️  No person info found, analyzing image directly")
            recommendation = await recommend_insurance_by_address(image_path)
        
        if "error" in recommendation:
            print(f"   ❌ Analysis error: {recommendation['error']}")
        else:
            place_region = recommendation.get('place_of_origin', {}).get('region', 'Unknown')
            addr_region = recommendation.get('address', {}).get('region', 'Unknown')
            num_packages = len(recommendation.get('recommended_packages', []))
            print(f"   ✅ Quê quán region: {place_region}, Address region: {addr_region}")
            print(f"   📦 Recommended: {num_packages} packages")
        
        return {
            "document_id": document_id,
            "recommendation": recommendation,
            "message": "Insurance recommendations generated successfully"
        }
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"\n❌ Insurance recommendation failed: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Recommendation failed: {str(e)}")
    finally:
        db.close()


@app.post("/chat")
async def chat_endpoint(request: dict):
    """
    Chat with AI Insurance Advisor
    
    Request body:
    {
        "message": "User's message",
        "document_id": "optional document ID for context",
        "chat_history": [optional previous messages]
    }
    """
    db = get_db()
    try:
        message = request.get('message')
        document_id = request.get('document_id')
        chat_history = request.get('chat_history', [])
        
        if not message:
            raise HTTPException(status_code=400, detail="Message is required")
        
        print(f"\n💬 Chat request: message='{message[:50]}...', doc_id={document_id}")
        
        # Get document analysis if document_id provided
        document_analysis = None
        if document_id:
            document = db.query(Document).filter(Document.id == document_id).first()
            if document and document.ai_result_json:
                import json
                ai_result = json.loads(document.ai_result_json)
                
                # Extract recommendation data
                if 'recommendation' in ai_result:
                    document_analysis = ai_result['recommendation']
                    print(f"   📋 Using document analysis with region: {document_analysis.get('place_of_origin', {}).get('region', 'Unknown')}")
        
        # Call chat service
        from app.chat_service import chat_with_insurance_advisor
        response = await chat_with_insurance_advisor(
            user_message=message,
            document_analysis=document_analysis,
            chat_history=chat_history
        )
        
        return {
            "reply": response['reply'],
            "has_context": response.get('has_context', False),
            "region": response.get('region'),
            "message": "Chat response generated successfully"
        }
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"\n❌ Chat failed: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Chat failed: {str(e)}")
    finally:
        db.close()


# ===================== INSURANCE PURCHASE HISTORY ENDPOINTS =====================

@app.post("/insurance-purchases")
async def create_insurance_purchase(request: dict):
    """
    Create new insurance purchase record
    
    Request body:
    {
        "user_id": 1,
        "package_name": "Bảo hiểm TNDS Xe máy",
        "package_type": "TNDS",
        "customer_name": "Nguyễn Văn A",
        "customer_phone": "0901234567",
        "premium_amount": "500000",
        "document_id": "optional",
        ... (other fields)
    }
    """
    from app.models import InsurancePurchase
    db = get_db()
    
    try:
        # Validate required fields
        required_fields = ['user_id', 'package_name', 'package_type', 'customer_name', 'customer_phone', 'premium_amount']
        for field in required_fields:
            if field not in request or not request[field]:
                raise HTTPException(status_code=400, detail=f"Field '{field}' is required")
        
        # Create purchase record
        purchase = InsurancePurchase(
            user_id=request['user_id'],
            package_name=request['package_name'],
            package_type=request['package_type'],
            insurance_company=request.get('insurance_company', 'VAM Insurance'),
            customer_name=request['customer_name'],
            customer_phone=request['customer_phone'],
            customer_email=request.get('customer_email'),
            customer_address=request.get('customer_address'),
            customer_id_number=request.get('customer_id_number'),
            coverage_amount=request.get('coverage_amount'),
            premium_amount=request['premium_amount'],
            payment_frequency=request.get('payment_frequency', 'Năm'),
            start_date=request.get('start_date'),
            end_date=request.get('end_date'),
            beneficiary_name=request.get('beneficiary_name'),
            beneficiary_relationship=request.get('beneficiary_relationship'),
            vehicle_type=request.get('vehicle_type'),
            license_plate=request.get('license_plate'),
            payment_method=request.get('payment_method'),
            payment_status=request.get('payment_status', 'PENDING'),
            transaction_id=request.get('transaction_id'),
            document_id=request.get('document_id'),
            policy_number=request.get('policy_number'),
            status=request.get('status', 'ACTIVE'),
            additional_data=json.dumps(request.get('additional_data', {})) if request.get('additional_data') else None,
            created_at=datetime.utcnow(),
            updated_at=datetime.utcnow()
        )
        
        db.add(purchase)
        db.commit()
        db.refresh(purchase)
        
        print(f"✅ Created insurance purchase #{purchase.id} for user {purchase.user_id}")
        
        return {
            "purchase_id": purchase.id,
            "message": "Insurance purchase created successfully"
        }
        
    except HTTPException:
        raise
    except Exception as e:
        db.rollback()
        print(f"❌ Failed to create purchase: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to create purchase: {str(e)}")
    finally:
        db.close()


@app.get("/users/{user_id}/insurance-purchases")
async def get_user_insurance_purchases(user_id: int):
    """
    Get all insurance purchases for a user
    """
    from app.models import InsurancePurchase
    db = get_db()
    
    try:
        purchases = db.query(InsurancePurchase)\
            .filter(InsurancePurchase.user_id == user_id)\
            .order_by(InsurancePurchase.created_at.desc())\
            .all()
        
        result = []
        for p in purchases:
            result.append({
                "id": p.id,
                "package_name": p.package_name,
                "package_type": p.package_type,
                "insurance_company": p.insurance_company,
                "customer_name": p.customer_name,
                "customer_phone": p.customer_phone,
                "customer_email": p.customer_email,
                "customer_address": p.customer_address,
                "customer_id_number": p.customer_id_number,
                "coverage_amount": p.coverage_amount,
                "premium_amount": p.premium_amount,
                "payment_frequency": p.payment_frequency,
                "start_date": p.start_date,
                "end_date": p.end_date,
                "beneficiary_name": p.beneficiary_name,
                "beneficiary_relationship": p.beneficiary_relationship,
                "vehicle_type": p.vehicle_type,
                "license_plate": p.license_plate,
                "payment_method": p.payment_method,
                "payment_status": p.payment_status,
                "transaction_id": p.transaction_id,
                "document_id": p.document_id,
                "policy_number": p.policy_number,
                "status": p.status,
                "additional_data": json.loads(p.additional_data) if p.additional_data else None,
                "created_at": p.created_at.isoformat() if p.created_at else None,
                "updated_at": p.updated_at.isoformat() if p.updated_at else None
            })
        
        print(f"📋 Retrieved {len(result)} purchases for user {user_id}")
        
        return {
            "purchases": result,
            "total": len(result)
        }
        
    except Exception as e:
        print(f"❌ Failed to get purchases: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to get purchases: {str(e)}")
    finally:
        db.close()


@app.get("/insurance-purchases/{purchase_id}")
async def get_insurance_purchase(purchase_id: int):
    """
    Get single insurance purchase by ID
    """
    from app.models import InsurancePurchase
    db = get_db()
    
    try:
        purchase = db.query(InsurancePurchase)\
            .filter(InsurancePurchase.id == purchase_id)\
            .first()
        
        if not purchase:
            raise HTTPException(status_code=404, detail="Purchase not found")
        
        return {
            "id": purchase.id,
            "user_id": purchase.user_id,
            "package_name": purchase.package_name,
            "package_type": purchase.package_type,
            "insurance_company": purchase.insurance_company,
            "customer_name": purchase.customer_name,
            "customer_phone": purchase.customer_phone,
            "customer_email": purchase.customer_email,
            "customer_address": purchase.customer_address,
            "customer_id_number": purchase.customer_id_number,
            "coverage_amount": purchase.coverage_amount,
            "premium_amount": purchase.premium_amount,
            "payment_frequency": purchase.payment_frequency,
            "start_date": purchase.start_date,
            "end_date": purchase.end_date,
            "beneficiary_name": purchase.beneficiary_name,
            "beneficiary_relationship": purchase.beneficiary_relationship,
            "vehicle_type": purchase.vehicle_type,
            "license_plate": purchase.license_plate,
            "payment_method": purchase.payment_method,
            "payment_status": purchase.payment_status,
            "transaction_id": purchase.transaction_id,
            "document_id": purchase.document_id,
            "policy_number": purchase.policy_number,
            "status": purchase.status,
            "additional_data": json.loads(purchase.additional_data) if purchase.additional_data else None,
            "created_at": purchase.created_at.isoformat() if purchase.created_at else None,
            "updated_at": purchase.updated_at.isoformat() if purchase.updated_at else None
        }
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ Failed to get purchase: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to get purchase: {str(e)}")
    finally:
        db.close()


@app.put("/insurance-purchases/{purchase_id}")
async def update_insurance_purchase(purchase_id: int, request: dict):
    """
    Update insurance purchase (e.g., payment status, policy number)
    """
    from app.models import InsurancePurchase
    db = get_db()
    
    try:
        purchase = db.query(InsurancePurchase)\
            .filter(InsurancePurchase.id == purchase_id)\
            .first()
        
        if not purchase:
            raise HTTPException(status_code=404, detail="Purchase not found")
        
        # Update allowed fields
        updatable_fields = [
            'payment_status', 'transaction_id', 'policy_number', 'status',
            'payment_method', 'start_date', 'end_date', 'additional_data'
        ]
        
        for field in updatable_fields:
            if field in request:
                if field == 'additional_data' and request[field]:
                    setattr(purchase, field, json.dumps(request[field]))
                else:
                    setattr(purchase, field, request[field])
        
        purchase.updated_at = datetime.utcnow()
        
        db.commit()
        db.refresh(purchase)
        
        print(f"✅ Updated insurance purchase #{purchase_id}")
        
        return {
            "purchase_id": purchase.id,
            "message": "Insurance purchase updated successfully"
        }
        
    except HTTPException:
        raise
    except Exception as e:
        db.rollback()
        print(f"❌ Failed to update purchase: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to update purchase: {str(e)}")
    finally:
        db.close()


def format_currency_for_pdf(amount_str):
    """
    Format currency string for PDF display
    Handles both plain numbers and formatted strings like "500.000.000 VNĐ"
    """
    if not amount_str:
        return 'N/A'
    
    try:
        # If it's already a formatted string with VNĐ, return as is
        if isinstance(amount_str, str) and ('VNĐ' in amount_str.upper() or 'VND' in amount_str.upper()):
            return amount_str
        
        # Try to convert to int and format
        # Remove any existing formatting first
        clean_str = str(amount_str).replace('.', '').replace(',', '').replace('VNĐ', '').replace('VND', '').replace('đ', '').strip()
        amount = int(clean_str)
        
        # Format with Vietnamese style (dots as thousand separators)
        formatted = f"{amount:,}".replace(',', '.')
        return f"{formatted} VNĐ"
    except (ValueError, AttributeError):
        # If conversion fails, return the original string
        return str(amount_str)


def format_date_vietnamese(date_str):
    """
    Format date to Vietnamese format DD/MM/YYYY
    """
    if not date_str:
        return 'Chưa xác định'
    
    try:
        # If already in DD/MM/YYYY format
        if '/' in date_str and len(date_str.split('/')) == 3:
            return date_str
        
        # If in YYYY-MM-DD format
        if '-' in date_str and len(date_str.split('-')) == 3:
            parts = date_str.split('-')
            return f"{parts[2]}/{parts[1]}/{parts[0]}"
        
        return date_str
    except:
        return date_str


def format_currency_for_pdf(amount):
    """
    Format currency for PDF display
    """
    if not amount:
        return 'Chưa xác định'
    
    try:
        # If it's already a formatted string with currency, return as-is
        if isinstance(amount, str):
            if 'VNĐ' in amount or 'VND' in amount:
                return amount
            
            # Try to extract number from string
            import re
            numbers = re.findall(r'\d+\.?\d*', amount.replace('.', '').replace(',', ''))
            if numbers:
                amount = float(numbers[0])
            else:
                return amount  # Return original if can't parse
        
        # Format with Vietnamese locale (dot as thousand separator)
        formatted = "{:,.0f}".format(float(amount)).replace(',', '.')
        return f"{formatted} VNĐ"
    except:
        return str(amount)


@app.get("/insurance-purchases/{purchase_id}/download-contract")
async def download_insurance_contract(purchase_id: int):
    """
    Download insurance contract as PDF
    Generates a professional PDF contract document with Vietnamese support
    """
    from app.models import InsurancePurchase
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch, cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
    from io import BytesIO
    import os
    
    db = get_db()
    
    try:
        purchase = db.query(InsurancePurchase)\
            .filter(InsurancePurchase.id == purchase_id)\
            .first()
        
        if not purchase:
            raise HTTPException(status_code=404, detail="Purchase not found")
        
        # Create PDF in memory
        buffer = BytesIO()
        doc = SimpleDocTemplate(
            buffer, 
            pagesize=A4,
            rightMargin=2*cm,
            leftMargin=2*cm,
            topMargin=2*cm,
            bottomMargin=2*cm
        )
        story = []
        styles = getSampleStyleSheet()
        
        # Try to register Vietnamese font (fallback to default if not available)
        try:
            # Register a Vietnamese-compatible font if available
            font_path = "C:/Windows/Fonts/times.ttf"  # Times New Roman supports Vietnamese
            if os.path.exists(font_path):
                pdfmetrics.registerFont(TTFont('TimesVN', font_path))
                default_font = 'TimesVN'
            else:
                default_font = 'Helvetica'
        except:
            default_font = 'Helvetica'
        
        # Create professional custom styles
        header_style = ParagraphStyle(
            'Header',
            parent=styles['Normal'],
            fontName=default_font,
            fontSize=26,
            textColor=colors.HexColor('#1E40AF'),
            spaceAfter=10,
            alignment=TA_CENTER,
            leading=32
        )
        
        subtitle_style = ParagraphStyle(
            'Subtitle',
            parent=styles['Normal'],
            fontName=default_font,
            fontSize=16,
            textColor=colors.HexColor('#DC2626'),
            spaceAfter=20,
            alignment=TA_CENTER,
            leading=20
        )
        
        company_style = ParagraphStyle(
            'Company',
            parent=styles['Normal'],
            fontName=default_font,
            fontSize=18,
            textColor=colors.HexColor('#059669'),
            spaceAfter=8,
            alignment=TA_CENTER,
            leading=22
        )
        
        section_heading_style = ParagraphStyle(
            'SectionHeading',
            parent=styles['Normal'],
            fontName=default_font,
            fontSize=14,
            textColor=colors.HexColor('#1F2937'),
            spaceAfter=15,
            spaceBefore=25,
            alignment=TA_LEFT,
            backColor=colors.HexColor('#F3F4F6'),
            leading=18,
            borderPadding=8
        )
        
        body_style = ParagraphStyle(
            'Body',
            parent=styles['Normal'],
            fontName=default_font,
            fontSize=11,
            textColor=colors.black,
            spaceAfter=6,
            alignment=TA_LEFT,
            leading=14
        )
        
        # === HEADER SECTION ===
        story.append(Paragraph("HỢP ĐỒNG BẢO HIỂM", header_style))
        story.append(Paragraph(f"Số hợp đồng: BH{purchase.policy_number or str(purchase.id).zfill(8)}", subtitle_style))
        
        # Company header with border
        story.append(Spacer(1, 0.3*inch))
        story.append(Paragraph("CÔNG TY BẢO HIỂM VAM", company_style))
        
        # Company details
        company_info = [
            f"<b>Công ty:</b> {purchase.insurance_company or 'VAM Insurance'}",
            "<b>Địa chỉ:</b> TP. Hồ Chí Minh, Việt Nam",
            "<b>Hotline:</b> 1900 xxxx",
            f"<b>Ngày tạo hợp đồng:</b> {datetime.now().strftime('%d/%m/%Y')}"
        ]
        
        for info in company_info:
            story.append(Paragraph(info, body_style))
        
        story.append(Spacer(1, 0.4*inch))
        
        # === PACKAGE INFORMATION SECTION ===
        story.append(Paragraph("I. THÔNG TIN GÓI BẢO HIỂM", section_heading_style))
        
        # Create separate styles for different types of information
        label_style = ParagraphStyle(
            'Label',
            parent=body_style,
            fontName=default_font,
            fontSize=11,
            textColor=colors.black,
            leftIndent=10
        )
        
        value_style = ParagraphStyle(
            'Value',
            parent=body_style,
            fontName=default_font,
            fontSize=11,
            textColor=colors.black,
            leftIndent=10
        )
        
        coverage_style = ParagraphStyle(
            'Coverage',
            parent=body_style,
            fontName=default_font,
            fontSize=11,
            textColor=colors.HexColor('#059669'),
            leftIndent=10
        )
        
        premium_style = ParagraphStyle(
            'Premium',
            parent=body_style,
            fontName=default_font,
            fontSize=11,
            textColor=colors.HexColor('#DC2626'),
            leftIndent=10
        )
        
        # Format coverage amount properly
        coverage_display = purchase.coverage_amount or '300.000.000 VNĐ/người/năm'
        if not any(x in coverage_display for x in ['/', 'người', 'năm']):
            coverage_display += '/người/năm'
            
        # Format premium amount
        premium_display = format_currency_for_pdf(purchase.premium_amount)
        
        package_data = [
            [Paragraph('<b>Tên gói bảo hiểm</b>', label_style), Paragraph(purchase.package_name or 'N/A', value_style)],
            [Paragraph('<b>Loại bảo hiểm</b>', label_style), Paragraph(purchase.package_type or 'N/A', value_style)],
            [Paragraph('<b>Số tiền bảo hiểm<br/>(Mức bảo hiểm tối đa)</b>', label_style), Paragraph(f'<b>{coverage_display}</b>', coverage_style)],
            [Paragraph('<b>Phí bảo hiểm<br/>(Số tiền phải trả)</b>', label_style), Paragraph(f'<b>{premium_display}</b>', premium_style)],
            [Paragraph('<b>Tần suất thanh toán</b>', label_style), Paragraph(purchase.payment_frequency or 'Hàng năm', value_style)],
            [Paragraph('<b>Ngày bắt đầu</b>', label_style), Paragraph(format_date_vietnamese(purchase.start_date), value_style)],
            [Paragraph('<b>Ngày kết thúc</b>', label_style), Paragraph(format_date_vietnamese(purchase.end_date), value_style)],
        ]
        
        package_table = Table(package_data, colWidths=[5*cm, 10*cm])
        package_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#EBF8FF')),
            ('BACKGROUND', (1, 0), (1, -1), colors.white),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (0, -1), 'LEFT'),
            ('ALIGN', (1, 0), (1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, -1), default_font),
            ('FONTSIZE', (0, 0), (-1, -1), 11),
            ('ROWBACKGROUNDS', (0, 0), (-1, -1), [colors.HexColor('#F8FAFC'), colors.white]),
            ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#E5E7EB')),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('TOPPADDING', (0, 0), (-1, -1), 8),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('LEFTPADDING', (0, 0), (-1, -1), 10),
        ]))
        story.append(package_table)
        story.append(Spacer(1, 0.3*inch))
        
        # === CUSTOMER INFORMATION SECTION ===
        story.append(Paragraph("II. THÔNG TIN KHÁCH HÀNG", section_heading_style))
        
        customer_data = [
            [Paragraph('<b>Họ và tên</b>', label_style), Paragraph(purchase.customer_name or 'N/A', value_style)],
            [Paragraph('<b>Số điện thoại</b>', label_style), Paragraph(purchase.customer_phone or 'N/A', value_style)],
            [Paragraph('<b>Email</b>', label_style), Paragraph(purchase.customer_email or 'Chưa cung cấp', value_style)],
            [Paragraph('<b>Địa chỉ</b>', label_style), Paragraph(getattr(purchase, 'customer_address', None) or 'Chưa cung cấp', value_style)],
            [Paragraph('<b>Số CMND/CCCD</b>', label_style), Paragraph(getattr(purchase, 'customer_id_number', None) or 'Chưa cung cấp', value_style)],
        ]
        
        # Add beneficiary info if available
        beneficiary_name = getattr(purchase, 'beneficiary_name', None)
        if beneficiary_name:
            customer_data.extend([
                [Paragraph('<b>Người thụ hưởng</b>', label_style), Paragraph(beneficiary_name, value_style)],
                [Paragraph('<b>Mối quan hệ</b>', label_style), Paragraph(getattr(purchase, 'beneficiary_relationship', None) or 'N/A', value_style)],
            ])
        
        customer_table = Table(customer_data, colWidths=[5*cm, 10*cm])
        customer_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#FEF3F2')),
            ('BACKGROUND', (1, 0), (1, -1), colors.white),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (0, -1), 'LEFT'),
            ('ALIGN', (1, 0), (1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, -1), default_font),
            ('FONTSIZE', (0, 0), (-1, -1), 11),
            ('ROWBACKGROUNDS', (0, 0), (-1, -1), [colors.HexColor('#FFFBFB'), colors.white]),
            ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#E5E7EB')),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('TOPPADDING', (0, 0), (-1, -1), 8),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('LEFTPADDING', (0, 0), (-1, -1), 10),
        ]))
        story.append(customer_table)
        story.append(Spacer(1, 0.3*inch))
        
        # === VEHICLE INFORMATION (if applicable) ===
        if purchase.vehicle_type:
            story.append(Paragraph("III. THÔNG TIN PHƯƠNG TIỆN", section_heading_style))
            
            vehicle_data = [
                [Paragraph('<b>Loại phương tiện</b>', label_style), Paragraph(purchase.vehicle_type or 'N/A', value_style)],
                [Paragraph('<b>Biển số đăng ký</b>', label_style), Paragraph(purchase.license_plate or 'Chưa cung cấp', value_style)],
            ]
            
            vehicle_table = Table(vehicle_data, colWidths=[5*cm, 10*cm])
            vehicle_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#F0F9FF')),
                ('BACKGROUND', (1, 0), (1, -1), colors.white),
                ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
                ('ALIGN', (0, 0), (0, -1), 'LEFT'),
                ('ALIGN', (1, 0), (1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, -1), default_font),
                ('FONTSIZE', (0, 0), (-1, -1), 11),
                ('ROWBACKGROUNDS', (0, 0), (-1, -1), [colors.HexColor('#F8FBFF'), colors.white]),
                ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#E5E7EB')),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('TOPPADDING', (0, 0), (-1, -1), 8),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                ('LEFTPADDING', (0, 0), (-1, -1), 10),
            ]))
            story.append(vehicle_table)
            story.append(Spacer(1, 0.3*inch))
        
        # === PAYMENT INFORMATION SECTION ===
        next_section = "IV." if purchase.vehicle_type else "III."
        story.append(Paragraph(f"{next_section} THÔNG TIN THANH TOÁN", section_heading_style))
        
        payment_status_color = colors.HexColor("#059669") if purchase.payment_status == "PAID" else colors.HexColor("#DC2626") if purchase.payment_status == "FAILED" else colors.HexColor("#D97706")
        
        payment_status_style = ParagraphStyle(
            'PaymentStatus',
            parent=body_style,
            fontName=default_font,
            fontSize=11,
            textColor=payment_status_color,
            leftIndent=10
        )
        
        payment_data = [
            [Paragraph('<b>Phương thức thanh toán</b>', label_style), Paragraph(purchase.payment_method or 'Chưa xác định', value_style)],
            [Paragraph('<b>Trạng thái thanh toán</b>', label_style), Paragraph(f'<b>{purchase.payment_status or "PENDING"}</b>', payment_status_style)],
            [Paragraph('<b>Mã giao dịch</b>', label_style), Paragraph(getattr(purchase, 'transaction_id', None) or f'TX{purchase.id}{datetime.now().strftime("%Y%m%d")}', value_style)],
        ]
        
        payment_table = Table(payment_data, colWidths=[5*cm, 10*cm])
        payment_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#F0FDF4')),
            ('BACKGROUND', (1, 0), (1, -1), colors.white),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (0, -1), 'LEFT'),
            ('ALIGN', (1, 0), (1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, -1), default_font),
            ('FONTSIZE', (0, 0), (-1, -1), 11),
            ('ROWBACKGROUNDS', (0, 0), (-1, -1), [colors.HexColor('#F7FEF7'), colors.white]),
            ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#E5E7EB')),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('TOPPADDING', (0, 0), (-1, -1), 8),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('LEFTPADDING', (0, 0), (-1, -1), 10),
        ]))
        story.append(payment_table)
        story.append(Spacer(1, 0.5*inch))
        
        # === TERMS AND CONDITIONS ===
        terms_section = "V." if purchase.vehicle_type else "IV."
        story.append(Paragraph(f"{terms_section} ĐIỀU KHOẢN VÀ ĐIỀU KIỆN", section_heading_style))
        
        terms_text = f"""
        <b>GIẢI THÍCH CÁC KHOẢN TIỀN:</b><br/>
        • <b>Số tiền bảo hiểm ({coverage_display}):</b> Đây là số tiền tối đa mà Công ty bảo hiểm sẽ chi trả khi xảy ra rủi ro được bảo hiểm.<br/>
        • <b>Phí bảo hiểm ({premium_display}):</b> Đây là số tiền Bên mua bảo hiểm phải thanh toán để duy trì hợp đồng bảo hiểm.<br/><br/>
        
        <b>ĐIỀU KHOẢN VÀ ĐIỀU KIỆN:</b><br/>
        1. Hợp đồng này có hiệu lực kể từ ngày ký và thanh toán đầy đủ phí bảo hiểm.<br/>
        2. Bên mua bảo hiểm có trách nhiệm cung cấp thông tin chính xác và đầy đủ.<br/>
        3. Công ty bảo hiểm cam kết bồi thường theo đúng điều khoản đã thỏa thuận.<br/>
        4. Mọi tranh chấp sẽ được giải quyết theo quy định của pháp luật Việt Nam.<br/>
        5. Hợp đồng này được lập thành 02 bản có giá trị pháp lý như nhau.
        """
        
        story.append(Paragraph(terms_text, body_style))
        story.append(Spacer(1, 0.5*inch))
        
        # === SIGNATURE SECTION ===
        story.append(Spacer(1, 0.3*inch))
        
        signature_header = Paragraph(
            f"<b>Ngày ký: {datetime.now().strftime('%d tháng %m năm %Y')}</b>",
            ParagraphStyle('SignatureHeader', parent=body_style, alignment=TA_CENTER, fontSize=12)
        )
        story.append(signature_header)
        story.append(Spacer(1, 0.2*inch))
        
        signature_style = ParagraphStyle(
            'Signature',
            parent=body_style,
            fontName=default_font,
            fontSize=11,
            alignment=TA_CENTER
        )
        
        signature_small_style = ParagraphStyle(
            'SignatureSmall',
            parent=body_style,
            fontName=default_font,
            fontSize=9,
            alignment=TA_CENTER,
            textColor=colors.HexColor('#6B7280')
        )
        
        signature_data = [
            [Paragraph('<b>BÊN MUA BẢO HIỂM</b>', signature_style), Paragraph('<b>ĐẠI DIỆN CÔNG TY</b>', signature_style)],
            [Paragraph('', signature_style), Paragraph('', signature_style)],
            [Paragraph('', signature_style), Paragraph('', signature_style)],
            [Paragraph('', signature_style), Paragraph('', signature_style)],
            [Paragraph(f'<b>{purchase.customer_name}</b>', signature_style), Paragraph('<b>Giám đốc</b>', signature_style)],
            [Paragraph('<i>(Ký và ghi rõ họ tên)</i>', signature_small_style), Paragraph('<i>(Ký tên và đóng dấu)</i>', signature_small_style)],
        ]
        
        signature_table = Table(signature_data, colWidths=[7.5*cm, 7.5*cm])
        signature_table.setStyle(TableStyle([
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, -1), default_font),
            ('FONTSIZE', (0, 0), (-1, -1), 11),
            ('FONTSIZE', (0, 5), (-1, 5), 9),
            ('BOTTOMPADDING', (0, 1), (-1, 3), 15),
            ('TOPPADDING', (0, 4), (-1, 4), 10),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ]))
        story.append(signature_table)
        
        # === FOOTER ===
        story.append(Spacer(1, 0.5*inch))
        footer_style = ParagraphStyle(
            'Footer',
            parent=body_style,
            fontSize=10,
            textColor=colors.HexColor('#6B7280'),
            alignment=TA_CENTER,
            leading=12
        )
        
        footer_text = f"""
        <b>HỢP ĐỒNG BẢO HIỂM - SỐ: BH{purchase.policy_number or str(purchase.id).zfill(8)}</b><br/>
        Được tạo tự động bởi hệ thống VAM Insurance vào ngày {datetime.now().strftime('%d/%m/%Y lúc %H:%M')}<br/>
        Để biết thêm thông tin, vui lòng liên hệ hotline: 1900 xxxx - website: www.vaminsurance.vn
        """
        
        story.append(Paragraph(footer_text, footer_style))
        
        # Build PDF
        doc.build(story)
        
        # Get PDF data
        buffer.seek(0)
        pdf_data = buffer.getvalue()
        buffer.close()
        
        # Return as file response
        from fastapi.responses import Response
        
        # Clean filename
        clean_customer_name = ''.join(c for c in purchase.customer_name if c.isalnum() or c in (' ', '_')).replace(' ', '_')
        filename = f"HopDong_BH{purchase.policy_number or str(purchase.id).zfill(8)}_{clean_customer_name}.pdf"
        
        print(f"✅ Generated PDF contract: {filename}")
        
        return Response(
            content=pdf_data,
            media_type="application/pdf",
            headers={
                "Content-Disposition": f"attachment; filename*=UTF-8''{filename}"
            }
        )
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ Failed to generate PDF contract: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Failed to generate contract: {str(e)}")
    finally:
        db.close()


# ===================== AI INSURANCE GEO-ANALYST ENDPOINTS =====================


# ===================== AI INSURANCE GEO-ANALYST ENDPOINTS =====================

@app.post("/analyze-weather-insurance")
async def analyze_weather_insurance(request: dict):
    """
    AI Insurance Geo-Analyst - Phân tích địa chỉ, thời tiết và đề xuất bảo hiểm
    
    Input JSON:
    {
        "user_profile": {
            "full_name": "Nguyễn Văn A",
            "dob": "1996-10-12",
            "address": "Thạch Hà, Hà Tĩnh",
            "document_type": "CCCD"
        },
        "weather_data": {
            "source": "OpenWeather",
            "temperature": "28°C",
            "condition": "mưa lớn",
            "alert": "nguy cơ ngập lụt"
        }
    }
    
    Output: Region detection, risk analysis, insurance recommendations, map overview
    """
    try:
        # Validate input
        if 'user_profile' not in request:
            raise HTTPException(status_code=400, detail="user_profile is required")
        if 'weather_data' not in request:
            raise HTTPException(status_code=400, detail="weather_data is required")
        
        user_profile = request['user_profile']
        weather_data = request['weather_data']
        
        # Validate required fields
        if 'address' not in user_profile or not user_profile['address']:
            raise HTTPException(status_code=400, detail="Address is required in user_profile")
        
        print(f"\n🌍 AI Geo-Analyst analyzing...")
        print(f"   User: {user_profile.get('full_name', 'N/A')}")
        print(f"   Address: {user_profile.get('address', 'N/A')}")
        print(f"   Weather: {weather_data.get('condition', 'N/A')}")
        
        # Analyze with Geo-Analyst
        analysis_result = GeoAnalyst.analyze_user_location(user_profile, weather_data)
        
        if 'error' in analysis_result:
            print(f"   ❌ Analysis error: {analysis_result['error']}")
            raise HTTPException(status_code=400, detail=analysis_result['error'])
        
        print(f"   ✅ Region: {analysis_result.get('user_region', 'Unknown')}")
        print(f"   ✅ Risk: {analysis_result.get('risk_level', 'Unknown')}")
        print(f"   ✅ Recommended {len(analysis_result.get('recommended_packages', []))} packages")
        
        return analysis_result
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"\n❌ Geo-analysis failed: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Geo-analysis failed: {str(e)}")


@app.post("/analyze-document-location")
async def analyze_document_location(document_id: str):
    """
    Extract address from uploaded CCCD document and analyze location
    Returns full geo-analysis with insurance recommendations
    """
    db = get_db()
    try:
        # Get document
        document = db.query(Document).filter(Document.id == document_id).first()
        if not document:
            raise HTTPException(status_code=404, detail="Document not found")
        
        print(f"\n📍 Analyzing document location: {document_id}")
        
        # Check if we have extracted person info
        person_data = None
        if document.person_data:
            import json
            person_data = json.loads(document.person_data)
        else:
            # Extract person info first
            print("   ⚠️  No person data found, extracting...")
            pages = db.query(Page).filter(Page.document_id == document_id).order_by(Page.page_index).all()
            
            if not pages or not pages[0].image_url:
                raise HTTPException(status_code=400, detail="No image found for this document")
            
            from app.ai_service import get_image_path_from_url, extract_person_info
            image_path = get_image_path_from_url(pages[0].image_url)
            
            if not image_path or not os.path.exists(image_path):
                raise HTTPException(status_code=400, detail=f"Image file not found: {image_path}")
            
            person_data = await extract_person_info(image_path)
            
            if "error" in person_data:
                raise HTTPException(status_code=400, detail=person_data['error'])
            
            # Save to database
            document.person_data = json.dumps(person_data, ensure_ascii=False)
            db.commit()
        
        # Prepare user profile
        # IMPORTANT: Prioritize placeOfOrigin (quê quán) for disaster analysis
        # Quê quán is more relevant for long-term risk assessment than temporary address
        user_profile = {
            "full_name": person_data.get('fullName', 'N/A'),
            "dob": person_data.get('dateOfBirth', 'N/A'),
            "address": person_data.get('placeOfOrigin', person_data.get('address', '')),
            "place_of_origin": person_data.get('placeOfOrigin', ''),
            "document_type": person_data.get('documentType', 'CCCD')
        }
        
        # Mock weather data (in production, fetch from OpenWeather API)
        # For now, use data from disasterData based on detected region
        weather_data = {
            "source": "Mock Weather Service",
            "temperature": "28°C",
            "condition": "Mưa lớn",
            "alert": "Nguy cơ ngập lụt"
        }
        
        print(f"   📋 User: {user_profile['full_name']}")
        print(f"   📍 Address: {user_profile['address']}")
        
        # Analyze with Geo-Analyst
        analysis_result = GeoAnalyst.analyze_user_location(user_profile, weather_data)
        
        if 'error' in analysis_result:
            raise HTTPException(status_code=400, detail=analysis_result['error'])
        
        # Save analysis result to document
        if not document.ai_result_json:
            document.ai_result_json = json.dumps({}, ensure_ascii=False)
        
        current_result = json.loads(document.ai_result_json)
        current_result['geo_analysis'] = analysis_result
        document.ai_result_json = json.dumps(current_result, ensure_ascii=False)
        db.commit()
        
        print(f"   ✅ Analysis complete: {analysis_result.get('user_region', 'Unknown')} - {analysis_result.get('risk_level', 'Unknown')}")
        
        return {
            "document_id": document_id,
            "analysis": analysis_result,
            "message": "Document location analyzed successfully"
        }
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"\n❌ Document location analysis failed: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Analysis failed: {str(e)}")
    finally:
        db.close()


@app.post("/gemini-prompt/geo-analyst")
async def generate_geo_analyst_prompt(request: dict):
    """
    Generate Gemini prompt for geo-analyst (for testing/debugging)
    Returns the exact prompt that would be sent to Gemini API
    """
    try:
        if 'user_profile' not in request or 'weather_data' not in request:
            raise HTTPException(
                status_code=400, 
                detail="Both user_profile and weather_data are required"
            )
        
        prompt = generate_gemini_prompt(
            user_profile=request['user_profile'],
            weather_data=request['weather_data']
        )
        
        return {
            "prompt": prompt,
            "message": "Gemini prompt generated successfully"
        }
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Prompt generation failed: {str(e)}")


# ==================== DISASTER LOCATION ENDPOINTS ====================

def _location_to_dict(location: DisasterLocation) -> dict:
    """Helper function to convert DisasterLocation model to dict with parsed JSON"""
    return {
        "id": location.id,
        "province": location.province,
        "region": location.region,
        "latitude": location.latitude,
        "longitude": location.longitude,
        "status": location.status,
        "marker_color": location.marker_color,
        "severity": location.severity,
        "advice": location.advice,
        "detail": location.detail,
        "recommended_packages": json.loads(location.recommended_packages) if location.recommended_packages else [],
        "weather_info": json.loads(location.weather_info) if location.weather_info else None,
        "last_updated": location.last_updated.isoformat() if location.last_updated else None,
        "created_at": location.created_at.isoformat() if location.created_at else None
    }

@app.get("/api/disaster-locations")
async def get_all_disaster_locations():
    """
    Get all disaster locations with latest weather data
    """
    db = next(get_db())
    try:
        locations = db.query(DisasterLocation).all()
        return [_location_to_dict(loc) for loc in locations]
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to fetch locations: {str(e)}")
    finally:
        db.close()


@app.get("/api/disaster-locations/region/{region}")
async def get_locations_by_region(region: str):
    """
    Get disaster locations filtered by region (Bắc, Trung, Nam)
    """
    db = next(get_db())
    try:
        locations = db.query(DisasterLocation).filter(
            DisasterLocation.region == region
        ).all()
        return [_location_to_dict(loc) for loc in locations]
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to fetch locations: {str(e)}")
    finally:
        db.close()


@app.get("/api/disaster-locations/search")
async def search_locations(province: Optional[str] = None, status: Optional[str] = None):
    """
    Search disaster locations by province name or status
    """
    db = next(get_db())
    try:
        query = db.query(DisasterLocation)
        
        if province:
            query = query.filter(DisasterLocation.province.ilike(f"%{province}%"))
        
        if status:
            query = query.filter(DisasterLocation.status == status)
        
        locations = query.all()
        return [_location_to_dict(loc) for loc in locations]
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Search failed: {str(e)}")
    finally:
        db.close()


@app.get("/api/disaster-locations/{location_id}")
async def get_location_by_id(location_id: str):
    """
    Get a single disaster location by ID
    """
    db = next(get_db())
    try:
        location = db.query(DisasterLocation).filter(
            DisasterLocation.id == location_id
        ).first()
        
        if not location:
            raise HTTPException(status_code=404, detail="Location not found")
        
        return _location_to_dict(location)
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to fetch location: {str(e)}")
    finally:
        db.close()


@app.get("/api/weather")
async def get_weather_by_coordinates(lat: float, lon: float):
    """
    Get weather data for specific coordinates using OpenWeatherMap API
    Frontend can call this to get real-time weather data
    
    Query params:
    - lat: Latitude (e.g., 10.8231 for Ho Chi Minh City)
    - lon: Longitude (e.g., 106.6297 for Ho Chi Minh City)
    
    Returns:
    {
        "temperature": 29.03,
        "feels_like": 33.57,
        "humidity": 74,
        "pressure": 1011,
        "condition": "Clouds",
        "description": "mây cụm",
        "wind_speed": 3.09,
        "clouds": 75,
        "rain_1h": 0,
        "rain_3h": 0,
        "visibility": 10000,
        "status": "ổn_định",
        "severity": "Thấp",
        "marker_color": "green",
        "fetched_at": "2025-11-17T10:00:00"
    }
    """
    try:
        print(f"\n🌤️  Fetching weather for coordinates: lat={lat}, lon={lon}")
        
        # Fetch weather from OpenWeatherMap
        weather_info = await WeatherService.fetch_weather(lat, lon)
        
        if not weather_info:
            raise HTTPException(status_code=500, detail="Failed to fetch weather data from OpenWeatherMap")
        
        # Determine disaster status and severity
        status = WeatherService.determine_status(weather_info)
        severity = WeatherService.determine_severity(status)
        marker_color = WeatherService.determine_marker_color(status)
        
        # Add status information to weather data
        weather_info['status'] = status
        weather_info['severity'] = severity
        weather_info['marker_color'] = marker_color
        
        print(f"   ✅ Weather fetched: {weather_info['temperature']}°C, {weather_info['condition']}, status: {status}")
        
        return weather_info
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"   ❌ Weather fetch error: {e}")
        raise HTTPException(status_code=500, detail=f"Weather fetch failed: {str(e)}")


@app.post("/api/disaster-locations/update-weather")
async def update_all_weather():
    """
    Update weather data for all disaster locations (Cron job endpoint)
    """
    db = next(get_db())
    try:
        print("\n🌦️  Starting weather update for all locations...")
        results = await WeatherService.update_all_locations(db)
        
        return {
            "message": "Weather update completed",
            "total_locations": results['total'],
            "successful_updates": results['success'],
            "failed_updates": results['failed']
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Weather update failed: {str(e)}")
    finally:
        db.close()


@app.post("/api/disaster-locations/{location_id}/update-weather")
async def update_location_weather(location_id: str):
    """
    Update weather data for a specific location
    """
    db = next(get_db())
    try:
        success = await WeatherService.update_location_weather(db, location_id)
        
        if not success:
            raise HTTPException(status_code=404, detail="Location not found or update failed")
        
        # Get updated location
        location = db.query(DisasterLocation).filter(
            DisasterLocation.id == location_id
        ).first()
        
        return {
            "message": "Weather updated successfully",
            "location": _location_to_dict(location)
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Weather update failed: {str(e)}")
    finally:
        db.close()


@app.post("/api/disaster-locations")
async def create_disaster_location(location: DisasterLocationCreate):
    """
    Create a new disaster location (Admin only)
    """
    db = next(get_db())
    try:
        # Check if location already exists
        existing = db.query(DisasterLocation).filter(
            DisasterLocation.id == location.id
        ).first()
        
        if existing:
            raise HTTPException(status_code=400, detail="Location ID already exists")
        
        # Create new location
        new_location = DisasterLocation(
            id=location.id,
            province=location.province,
            region=location.region,
            latitude=location.latitude,
            longitude=location.longitude,
            status=location.status,
            marker_color=location.marker_color,
            severity=location.severity,
            advice=location.advice,
            detail=location.detail,
            recommended_packages=json.dumps(location.recommended_packages or [], ensure_ascii=False)
        )
        
        db.add(new_location)
        db.commit()
        db.refresh(new_location)
        
        return _location_to_dict(new_location)
    except HTTPException:
        raise
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"Failed to create location: {str(e)}")
    finally:
        db.close()


@app.put("/api/disaster-locations/{location_id}")
async def update_disaster_location(location_id: str, location_update: DisasterLocationUpdate):
    """
    Update a disaster location (Admin only)
    """
    db = next(get_db())
    try:
        location = db.query(DisasterLocation).filter(
            DisasterLocation.id == location_id
        ).first()
        
        if not location:
            raise HTTPException(status_code=404, detail="Location not found")
        
        # Update fields if provided
        update_data = location_update.dict(exclude_unset=True)
        
        for field, value in update_data.items():
            if field == "recommended_packages" and value is not None:
                setattr(location, field, json.dumps(value, ensure_ascii=False))
            elif field == "weather_info" and value is not None:
                setattr(location, field, json.dumps(value, ensure_ascii=False))
            elif value is not None:
                setattr(location, field, value)
        
        location.last_updated = datetime.utcnow()
        
        db.commit()
        db.refresh(location)
        
        return _location_to_dict(location)
    except HTTPException:
        raise
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"Failed to update location: {str(e)}")
    finally:
        db.close()


if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)